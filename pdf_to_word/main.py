# app.py - Exact Text Extractor from Images & PDFs (Gemini AI)
# Works with older AND newer google-generativeai SDKs

import os
import io
import re
import tempfile
import streamlit as st
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image
from docx import Document
from docx.shared import Pt

# Load environment variables
load_dotenv()

# Configure Gemini API
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    st.error("GOOGLE_API_KEY not found. Please set it in your .env file or Streamlit Secrets.")
else:
    genai.configure(api_key=API_KEY)

# Initialize Streamlit UI
st.set_page_config(page_title="📄 Exact Text Extractor (Images & PDFs)", layout="centered")
st.title("📄 Exact Text Extractor from Images & PDFs")
st.write("""
Upload one or more **images (JPG/PNG)** or **PDF files**.  
The app will extract text **exactly as it appears** — no changes, no summaries.
""")
st.markdown("---")

# File uploader - supports both images and PDFs
uploaded_files = st.file_uploader(
    "Choose image or PDF files...",
    type=["jpg", "jpeg", "png", "pdf"],
    accept_multiple_files=True
)

# Global storage
extracted_data = []

if uploaded_files:
    # Sort by number in filename: 1.jpg, 2.pdf, etc.
    def extract_number(file):
        match = re.search(r'(\d+)', file.name)
        return int(match.group(1)) if match else float('inf')

    sorted_files = sorted(uploaded_files, key=extract_number)

    st.write("📁 Processing order:")
    for i, f in enumerate(sorted_files):
        num = re.search(r'(\d+)', f.name)
        q = f"Q{num.group(1)}" if num else "Unknown"
        st.write(f"{i+1}. `{f.name}` → {q}")
    st.markdown("---")

    if st.button("🔍 Extract TEXT EXACTLY As Shown"):
        if not API_KEY:
            st.warning("Set GOOGLE_API_KEY in .env or secrets.")
            st.stop()

        try:
            model = genai.GenerativeModel('gemini-2.5-flash')
        except Exception as e:
            st.error(f"Failed to initialize model: {e}")
            st.stop()

        st.info("Extracting text **exactly as written** in files...")
        extracted_data.clear()

        for file in sorted_files:
            with st.spinner(f"Processing: {file.name}"):
                try:
                    file_ext = file.name.lower().split('.')[-1]
                    full_text = ""

                    if file_ext in ['jpg', 'jpeg', 'png']:
                        # === Handle IMAGE ===
                        image_bytes = file.read()
                        image_stream = io.BytesIO(image_bytes)
                        image = Image.open(image_stream)

                        prompt = '''
                        You are an OCR tool. Your job is to transcribe ALL visible text from the image **EXACTLY as it appears**, without any changes.

                        Rules:
                        1. Do NOT summarize, interpret, or explain.
                        2. Do NOT fix grammar or spelling.
                        3. Preserve ALL original formatting:
                           - Line breaks
                           - Punctuation
                           - Spacing
                           - Bilingual text (Nepali + English together)
                           - Numbers, symbols, brackets
                        4. If text is cut off, write [text unclear] at that point.
                        5. Maintain the same visual order (top to bottom, left to right).
                        6. Do NOT add your own headings or labels.

                        Output only the raw, unmodified text exactly as seen.
                        '''

                        response = model.generate_content([prompt, image])
                        full_text = response.text.strip() if response.text else "[No text detected or low confidence]"

                    elif file_ext == 'pdf':
                        # === Handle PDF (compatible with older SDK) ===
                        pdf_bytes = file.read()

                        # Write to temporary file
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                            tmp.write(pdf_bytes)
                            tmp_path = tmp.name

                        try:
                            uploaded_file_obj = genai.upload_file(
                                path=tmp_path,
                                display_name=file.name,
                                mime_type="application/pdf"
                            )

                            prompt = '''
                            You are an OCR tool. Your job is to transcribe ALL visible text from the PDF **EXACTLY as it appears**, without any changes.

                            Rules:
                            1. Do NOT summarize, interpret, or explain.
                            2. Do NOT fix grammar or spelling.
                            3. Preserve ALL original formatting:
                               - Line breaks
                               - Punctuation
                               - Spacing
                               - Bilingual text (Nepali + English together)
                               - Numbers, symbols, brackets
                            4. If text is cut off, write [text unclear] at that point.
                            5. Maintain the same visual order (top to bottom, left to right).
                            6. Do NOT add your own headings or labels.

                            Output only the raw, unmodified text exactly as seen.
                            '''

                            response = model.generate_content([prompt, uploaded_file_obj])
                            full_text = response.text.strip() if response.text else "[No text detected or low confidence]"

                        except Exception as e:
                            full_text = f"[Gemini error: {str(e)}]"
                        finally:
                            # Clean up temp file
                            if os.path.exists(tmp_path):
                                os.unlink(tmp_path)

                    else:
                        full_text = "[Unsupported file type]"

                    # Optional: Clean minor artifacts (keep empty lines for structure)
                    lines = [line.strip() for line in full_text.split('\n')]
                    cleaned_lines = []
                    for line in lines:
                        if line == '':
                            cleaned_lines.append('')
                        else:
                            cleaned_lines.append(line)
                    final_text = '\n'.join(cleaned_lines)

                    # Store and display
                    extracted_data.append({
                        "filename": file.name,
                        "text": final_text
                    })

                    st.markdown(f"### 📄 Extracted from: `{file.name}`")
                    st.text_area("", final_text, height=300, key=f"text_{file.name}")
                    st.markdown("<hr>", unsafe_allow_html=True)

                except Exception as e:
                    error_msg = f"[Unexpected error: {str(e)}]"
                    extracted_data.append({"filename": file.name, "text": error_msg})
                    st.error(error_msg)
                    st.markdown("<hr>", unsafe_allow_html=True)

        st.success("✅ Extraction complete! Ready to download.")

        # --- Generate .docx ---
        doc = Document()
        doc.add_heading('Exact Text Extraction from Images & PDFs', 0)
        doc.add_paragraph("Files processed in order shown above.")

        for item in extracted_data:
            # Filename header
            p = doc.add_paragraph()
            run = p.add_run(f"Source: {item['filename']}")
            run.bold = True
            run.font.size = Pt(12)

            # Add text line by line to preserve formatting
            text = item["text"]
            for line in text.split("\n"):
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(11)
                # Optional: set font for Nepali support
                run.font.name = 'Nirmala UI'

            # Separator
            doc.add_paragraph("_" * 60)

        # Save to buffer
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)

        # Download button
        st.download_button(
            label="⬇️ Download All as Word (.docx)",
            data=bio,
            file_name="exact_text_extraction.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Optional: Show SDK version for debugging
# st.caption(f"google-generativeai v{genai.__version__}")